import os
import logging
import tempfile
import json
from datetime import datetime
import PyPDF2
import docx
from app import db
from models import Document, MemoryEntry
import google_services

# Configure logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

def process_document(user, document_data, document_type=None, memory_id=None):
    """Process a document and extract information."""
    try:
        # Create document entry in database
        document = Document(
            user_id=user.id,
            title=document_data.get('name', 'Untitled Document'),
            file_type=document_type or document_data.get('mimeType', 'unknown'),
            created_at=datetime.utcnow()
        )
        
        if 'id' in document_data:
            document.drive_id = document_data['id']
        
        # Extract text content based on file type
        content_text = extract_text_from_document(user, document)
        if content_text:
            document.content_text = content_text
        
        # Link to memory if provided
        if memory_id:
            document.memory_id = memory_id
        
        db.session.add(document)
        db.session.commit()
        
        return document
    except Exception as e:
        logger.error(f"Error processing document: {e}")
        db.session.rollback()
        return None

def extract_text_from_document(user, document):
    """Extract text content from a document based on its type."""
    try:
        if not document.drive_id:
            return None
        
        # Get Drive service
        drive_service = google_services.get_drive_service(user)
        if not drive_service:
            return None
        
        # Download file content
        from googleapiclient.http import MediaIoBaseDownload
        request = drive_service.files().get_media(fileId=document.drive_id)
        
        with tempfile.NamedTemporaryFile(delete=False) as temp_file:
            downloader = MediaIoBaseDownload(temp_file, request)
            done = False
            while not done:
                status, done = downloader.next_chunk()
            
            temp_file_path = temp_file.name
        
        # Extract text based on file type
        extracted_text = ""
        
        try:
            mime_type = document.file_type.lower()
            
            if 'pdf' in mime_type:
                # Process PDF
                with open(temp_file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page_num in range(len(pdf_reader.pages)):
                        page = pdf_reader.pages[page_num]
                        extracted_text += page.extract_text() + "\n"
            
            elif 'word' in mime_type or 'docx' in mime_type:
                # Process Word document
                doc = docx.Document(temp_file_path)
                for para in doc.paragraphs:
                    extracted_text += para.text + "\n"
            
            elif 'text' in mime_type or 'txt' in mime_type:
                # Process plain text
                with open(temp_file_path, 'r', encoding='utf-8', errors='ignore') as file:
                    extracted_text = file.read()
            
            else:
                logger.warning(f"Unsupported file type for text extraction: {mime_type}")
        finally:
            # Clean up temporary file
            if os.path.exists(temp_file_path):
                os.unlink(temp_file_path)
        
        return extracted_text if extracted_text else None
    except Exception as e:
        logger.error(f"Error extracting text from document: {e}")
        return None

def create_document_summary(user, document_id):
    """Create a summary of a document using OpenManus."""
    try:
        document = Document.query.get(document_id)
        if not document or document.user_id != user.id:
            return "Document not found or access denied"
        
        if not document.content_text:
            return "No text content available to summarize"
        
        # Use OpenManus to generate summary
        from manus_integration import generate_document_summary
        
        summary = generate_document_summary(document.content_text)
        
        # Create memory entry for the summary
        memory_entry = MemoryEntry(
            user_id=user.id,
            entry_type='document_summary',
            title=f"Summary of {document.title}",
            content=summary,
            created_at=datetime.utcnow()
        )
        
        # Link document to this memory
        document.memory_id = memory_entry.id
        
        db.session.add(memory_entry)
        db.session.commit()
        
        return summary
    except Exception as e:
        logger.error(f"Error creating document summary: {e}")
        db.session.rollback()
        return f"Error creating summary: {str(e)}"

def categorize_document(user, document_id, category):
    """Categorize a document and file it appropriately."""
    try:
        document = Document.query.get(document_id)
        if not document or document.user_id != user.id:
            return "Document not found or access denied"
        
        # Create or update memory entry for this document
        if document.memory_id:
            memory_entry = MemoryEntry.query.get(document.memory_id)
            if memory_entry:
                # Update existing memory
                current_metadata = memory_entry.metadata or {}
                current_metadata['category'] = category
                memory_entry.metadata = current_metadata
                memory_entry.updated_at = datetime.utcnow()
            else:
                # Create new memory if linked memory doesn't exist
                memory_entry = MemoryEntry(
                    user_id=user.id,
                    entry_type='document',
                    title=document.title,
                    content=f"Document: {document.title}",
                    metadata={'category': category},
                    created_at=datetime.utcnow()
                )
                db.session.add(memory_entry)
                db.session.commit()
                document.memory_id = memory_entry.id
        else:
            # Create new memory entry
            memory_entry = MemoryEntry(
                user_id=user.id,
                entry_type='document',
                title=document.title,
                content=f"Document: {document.title}",
                metadata={'category': category},
                created_at=datetime.utcnow()
            )
            db.session.add(memory_entry)
            db.session.commit()
            document.memory_id = memory_entry.id
        
        # If the document is in Google Drive, we can organize it there too
        if document.drive_id:
            # This would involve moving the file to a folder with the category name
            # or applying labels/properties to the file
            drive_service = google_services.get_drive_service(user)
            if drive_service:
                # This is a simplified implementation - in a real app, you'd check if the folder exists
                # and create it if not, then move the file to that folder
                try:
                    # Find or create a folder with the category name
                    folder_query = f"mimeType='application/vnd.google-apps.folder' and name='{category}'"
                    folder_results = drive_service.files().list(
                        q=folder_query, pageSize=1, fields="files(id, name)"
                    ).execute()
                    
                    folder_id = None
                    folders = folder_results.get('files', [])
                    
                    if folders:
                        folder_id = folders[0]['id']
                    else:
                        # Create a new folder
                        folder_metadata = {
                            'name': category,
                            'mimeType': 'application/vnd.google-apps.folder'
                        }
                        folder = drive_service.files().create(
                            body=folder_metadata, fields='id'
                        ).execute()
                        folder_id = folder.get('id')
                    
                    if folder_id:
                        # Move the file to the folder
                        drive_service.files().update(
                            fileId=document.drive_id,
                            addParents=folder_id,
                            fields='id, parents'
                        ).execute()
                
                except Exception as e:
                    logger.error(f"Error organizing file in Drive: {e}")
        
        db.session.commit()
        return f"Document categorized as '{category}'"
    except Exception as e:
        logger.error(f"Error categorizing document: {e}")
        db.session.rollback()
        return f"Error categorizing document: {str(e)}"

def search_documents(user, query, category=None, limit=10):
    """Search documents based on content and metadata."""
    try:
        # Base query
        document_query = Document.query.filter_by(user_id=user.id)
        
        # Filter by category if specified
        if category:
            # Join with memory entries and filter by category in metadata
            document_query = document_query.join(
                MemoryEntry, Document.memory_id == MemoryEntry.id
            ).filter(
                MemoryEntry.metadata.contains({'category': category})
            )
        
        # Filter by content if query specified
        if query:
            document_query = document_query.filter(Document.content_text.ilike(f"%{query}%"))
        
        # Get results
        documents = document_query.limit(limit).all()
        
        result_list = []
        for doc in documents:
            result = {
                'id': doc.id,
                'title': doc.title,
                'file_type': doc.file_type,
                'created_at': doc.created_at.isoformat(),
                'drive_id': doc.drive_id
            }
            
            # Get category from memory if available
            if doc.memory_id:
                memory = MemoryEntry.query.get(doc.memory_id)
                if memory and memory.metadata and 'category' in memory.metadata:
                    result['category'] = memory.metadata['category']
            
            result_list.append(result)
        
        return result_list
    except Exception as e:
        logger.error(f"Error searching documents: {e}")
        return []
